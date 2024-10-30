import os
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
import csv

def convert_document(input_path, output_format):
    filename = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"uploads/{filename}.{output_format}"

    if output_format == 'pdf':
        doc = Document(input_path)
        doc.save(output_path)
    elif output_format == 'docx':
        if input_path.endswith('.pdf'):
            doc = Document()
            pdf = PdfReader(input_path)
            for page in pdf.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(output_path)
        else:
            with open(input_path, 'r') as file:
                content = file.read()
            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)
    elif output_format == 'txt':
        if input_path.endswith('.pdf'):
            pdf = PdfReader(input_path)
            with open(output_path, 'w') as file:
                for page in pdf.pages:
                    file.write(page.extract_text())
        else:
            doc = Document(input_path)
            with open(output_path, 'w') as file:
                for para in doc.paragraphs:
                    file.write(para.text + '\n')
    elif output_format == 'csv':
        doc = Document(input_path)
        with open(output_path, 'w', newline='') as file:
            writer = csv.writer(file)
            for para in doc.paragraphs:
                writer.writerow([para.text])

    return output_path